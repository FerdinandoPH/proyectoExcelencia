#Open a .docx file and only print the text in blue
import docx,sys,time
figs=[]
orderedfigs=[]
if len(sys.argv) < 2:
    doc=docx.Document(input('Pon el nombre del archivo Word: '))
else:
    doc = docx.Document(sys.argv[1])
for para in doc.paragraphs:
    for run in para.runs:
        if run.font.color.rgb == docx.shared.RGBColor(0,112,192):
            figs.append(run.text)
print(figs)
time.sleep(1)
figura=""
for element in figs:
    if "Fig" in element and figs.index(element) != 0:
        if figura[3]!='.':
            figura=figura[:3]+'.'+figura[3:]
        orderedfigs.append(figura)
        figura=element
    else:
        figura+=element
for orderedfig in orderedfigs:
    print(orderedfig)
